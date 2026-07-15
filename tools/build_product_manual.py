from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "artifacts" / "product-manual" / "StoneOS_Product_Manual.docx"
SCREENS = ROOT / "artifacts" / "product-manual" / "screens"

INK = "16201B"
PAPER = "F3F0E7"
BRASS = "B48A43"
MOSS = "55705D"
RUST = "A4553D"
MUTED = "6F776F"
LINE = "D8D1C1"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_geometry(table, widths_dxa):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_text(doc, text, bold_prefix=None, after=6, size=11, color=INK, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        set_font(first, size=size, color=color, bold=True)
        rest = p.add_run(text[len(bold_prefix):])
        set_font(rest, size=size, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, size=size, color=color, italic=italic)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_font(run)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_font(run)


def add_callout(doc, label, text, fill="E8EEF5", accent=MOSS):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.2
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "single")
    start.set(qn("w:sz"), "28")
    start.set(qn("w:space"), "8")
    start.set(qn("w:color"), accent)
    borders.append(start)
    p_pr.append(borders)
    label_run = p.add_run(f"{label}: ")
    set_font(label_run, color=accent, bold=True)
    text_run = p.add_run(text)
    set_font(text_run)


def add_table(doc, headers, rows, widths_dxa):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, "E8EEF5")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_font(run, size=9.5, color=DARK_BLUE, bold=True)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            p = cells[index].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_font(run, size=9.5, bold=index == 0)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1, page_break=False):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.page_break_before = page_break
    p.paragraph_format.keep_with_next = True
    return p


def add_screen(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    shape = run.add_picture(str(SCREENS / filename), width=Inches(6.45))
    doc_pr = shape._inline.docPr
    doc_pr.set("descr", caption)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    cr = cp.add_run(caption)
    set_font(cr, size=8.5, color=MUTED, italic=True)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)

for list_name in ("List Bullet", "List Number"):
    style = styles[list_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("StoneOS Product Manual  |  User and Team Edition")
set_font(hr, size=9, color=MUTED, bold=True)
footer = section.footer
add_page_number(footer.paragraphs[0])

# Cover - editorial_cover pattern with StoneOS palette.
doc.add_paragraph().paragraph_format.space_after = Pt(58)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kicker.paragraph_format.space_after = Pt(12)
kr = kicker.add_run("PRODUCT MANUAL · USER AND TEAM EDITION")
set_font(kr, size=11, color=BRASS, bold=True)
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
tr = title.add_run("StoneOS")
set_font(tr, size=34, color=INK, bold=True)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(24)
sr = subtitle.add_run("From raw block to revenue — one trusted factory workflow")
set_font(sr, size=15, color=MOSS)
add_screen(doc, "01-control-room.png", "StoneOS Control Room shown with fictional Atlas Stone Works data.")
add_callout(doc, "Demonstration-data notice", "Every factory name, serial number, customer, quantity and financial value in this manual and its companion video is fictional. Replace the examples with your own approved data during training.", fill="FFF6D9", accent=BRASS)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Edition 1.0 · July 2026 · Internal enablement and customer onboarding")
set_font(mr, size=9.5, color=MUTED, italic=True)

# 1. Orientation
add_heading(doc, "1. What StoneOS does", page_break=True)
add_text(doc, "StoneOS is the operating system for a stone factory. It records each physical event once and carries the same material identity from receiving through production, inventory, reservation, dispatch, invoicing and payment.")
add_callout(doc, "Core workflow", "Receive → Cut → Grind → Apply epoxy → Polish → Reserve → Dispatch → Bill and record payment.")
add_heading(doc, "The operating principles", 2)
add_bullets(doc, [
    "One operational truth: current stage, location and reservation status must agree.",
    "Traceability: every slab retains its parent block and production-session history.",
    "Role clarity: operators, supervisors, managers and owners see the work relevant to them.",
    "Append-only history: corrections are made through controlled reversals instead of deleting events.",
    "Tenant isolation: every record and query belongs to one factory.",
])
add_heading(doc, "Recommended daily sequence", 2)
add_steps(doc, [
    "Open My Work and review the live factory pulse.",
    "Record the next physical event where it happens.",
    "Resolve eligibility or reservation warnings before moving material.",
    "Review end-of-shift production, inventory and dispatch exceptions.",
])
add_heading(doc, "Navigation map", 2)
add_table(doc, ["Workspace", "Primary job", "Typical users"], [
    ["My Work", "Live counts and next actions", "All roles"],
    ["Receive / Inventory", "Goods receipt, on-hand stock and movements", "Inventory, supervisor, manager"],
    ["Cutting / Polishing", "B-21 and LPM production", "Operator, supervisor"],
    ["Orders & Dispatch", "Reservation, delivery, invoice and payment", "Sales, manager, owner"],
    ["Administration", "Opening setup, machines, imports and team access", "Manager, owner"],
], [2200, 4460, 2700])

# 2. Control room
add_heading(doc, "2. My Work: the Control Room", page_break=True)
add_screen(doc, "01-control-room.png", "Mock Control Room showing live operational pressure and a recommended next action.")
add_text(doc, "Use the Control Room at the beginning of a shift and before changing priorities. The cards summarize raw blocks, unpolished slabs, finished stock, reservations and current commercial activity.")
add_heading(doc, "How to read it", 2)
add_bullets(doc, [
    "Workflow rail: shows where material is concentrated from Receive through Bill.",
    "Operational pulse: compares current counts without waiting for end-of-day consolidation.",
    "Factory flow map: highlights pressure at yard, machines and stock locations.",
    "Next best action: suggests the next operational move using current workflow state.",
])
add_callout(doc, "Good practice", "Treat recommendations as decision support. The supervisor or manager remains responsible for machine availability, quality observations and safe physical handling.", fill="E7EFE8")

# 3. Receiving
add_heading(doc, "3. Receive raw blocks", page_break=True)
add_screen(doc, "02-receive.png", "Mock raw-block receipt for fictional block V104.")
add_heading(doc, "Procedure", 2)
add_steps(doc, [
    "Open Receive and confirm the physical receipt date.",
    "Enter the block serial, variety, weight and Raw Yard location.",
    "Add additional lines when several blocks arrive on the same receipt.",
    "Review the receipt lines, then record the receipt.",
])
add_callout(doc, "Result", "The block becomes available stock in the Raw Yard and a goods-receipt movement is appended with the operator and time. Supplier invoices may be entered later.")
add_heading(doc, "Before you save", 2)
add_bullets(doc, ["Serial number matches the physical marking.", "Weight unit and decimal are correct.", "Factory location reflects where the block was placed.", "No customer, supplier or financial information is invented to complete the receipt."])

# 4. Inventory
add_heading(doc, "4. Inventory and material trace", page_break=True)
add_screen(doc, "05-inventory.png", "Mock inventory screen tracing slab V101/50/01 through production.")
add_text(doc, "Inventory is the current operational state of each block and slab. Use filters and movement history to answer where material is, why it is there and what may happen next.")
add_heading(doc, "Common states", 2)
add_table(doc, ["State", "Meaning", "Expected location"], [
    ["Ready for cutting", "Raw block available for B-21 allocation", "Raw Yard or B-21 Queue"],
    ["Awaiting grinding", "Good slab created from a completed cut", "Unpolished Stock"],
    ["Grinding complete", "First LPM pass finished", "LPM Queue"],
    ["Epoxy applied", "Epoxy gate confirmed", "LPM Queue"],
    ["Sale ready", "Final polishing complete", "Finished Stock"],
    ["Reserved", "Held for a customer order", "Current stock location until dispatch"],
], [2100, 4680, 2580])
add_callout(doc, "Manager exceptions", "Use an adjustment only when the physical count differs from StoneOS. Enter a clear reason. Reverse the exact movement if the correction itself is wrong.", fill="FFF2ED", accent=RUST)

# 5. Cutting
add_heading(doc, "5. B-21 cutting", page_break=True)
add_screen(doc, "03-cutting.png", "Mock B-21 session for fictional raw block V101.")
add_heading(doc, "Start and log the session", 2)
add_steps(doc, [
    "Choose an eligible raw block and the B-21 machine.",
    "Start the cutting session; StoneOS reserves the block and moves it to B-21 WIP.",
    "Record operational-day runtime, downtime, power and notes as required.",
    "At completion, enter total slabs cut and final good slabs after inspection.",
    "Enter the shared length and width for the good batch, then complete the session.",
])
add_callout(doc, "Example", "50 slabs cut minus 3 damaged slabs produces 47 inventory slabs. StoneOS generates V101/50/01 through V101/50/47 and links them to block V101.")
add_heading(doc, "Important rules", 2)
add_bullets(doc, ["Damaged pieces do not become slab inventory.", "Do not complete or abort a session after moving or adjusting its reserved material; reconcile the inventory warning first.", "The daily production summary is derived from production records and should not be re-entered separately."])

# 6. LPM
add_heading(doc, "6. LPM finishing: grinding, epoxy and polishing", page_break=True)
add_screen(doc, "04-lpm.png", "Mock LPM screen showing the epoxy gate between grinding and polishing.")
add_text(doc, "StoneOS treats the LPM workflow as three distinct events. The next step is available only when the slab stage, location and reservation state match.")
add_table(doc, ["Step", "Eligible input", "Completion result"], [
    ["Grinding", "Awaiting-grinding slab in eligible stock", "Grinding Complete in LPM Queue"],
    ["Epoxy", "Ground, available slab in LPM Queue", "Epoxy Applied in LPM Queue"],
    ["Polishing", "Epoxy-applied slab in LPM Queue", "Sale Ready in Finished Stock"],
], [1700, 3830, 3830])
add_heading(doc, "Operating procedure", 2)
add_steps(doc, [
    "Select Grinding, choose the LPM machine and eligible slabs, then start the run.",
    "Complete the run only after the physical grinding pass is finished.",
    "Mark epoxy applied after the coating step is physically complete.",
    "Switch to Polishing, select only epoxy-applied slabs and start the final pass.",
    "Complete polishing to move slabs into Finished Stock as sale-ready inventory.",
])
add_callout(doc, "Correction rule", "Epoxy may be reversed only while the slab is still available in the epoxy-applied LPM Queue. Once polishing starts, the reversal is blocked.", fill="FFF6D9", accent=BRASS)

# 7. Sales
add_heading(doc, "7. Orders, reservation, dispatch and billing", page_break=True)
add_screen(doc, "06-sales.png", "Mock sales workflow for fictional customer Northstar Surfaces.")
add_steps(doc, [
    "Create or select the customer, then create a sales order.",
    "Select sale-ready slabs and reserve them for the customer.",
    "At dispatch, select the slabs actually loaded; unselected slabs remain reserved.",
    "Create the commercial invoice after reservation or delivery according to policy.",
    "Record payments against the invoice and monitor the remaining balance.",
])
add_callout(doc, "Traceability", "The reservation, delivery, invoice and payment remain linked to the same slab identities used during production.")
add_heading(doc, "Sales checks", 2)
add_bullets(doc, ["Confirm actual sale-time square feet on each line.", "Confirm vehicle and dispatch date before posting delivery.", "Use partial dispatch rather than releasing slabs that will ship later.", "Keep historical imports in Historical Data, not in the daily Sales workflow."])

# 8. Roles
add_heading(doc, "8. Roles, access and accountability", page_break=True)
add_screen(doc, "07-roles.png", "Mock role workspaces illustrating least-clutter access.")
add_table(doc, ["Role", "Primary responsibility", "Key access"], [
    ["Operator", "Record production as it happens", "Cutting and LPM"],
    ["Supervisor", "Keep the physical flow moving", "Receive, Inventory, Cutting, LPM, Sales visibility"],
    ["Inventory", "Receive and control stock movement", "Receive, Inventory, Orders & Dispatch"],
    ["Sales", "Reserve and dispatch finished slabs", "Inventory and Orders & Dispatch"],
    ["Manager", "Control operations and exceptions", "Operations, Expenses, Administration"],
    ["Owner", "Final authority and user provisioning", "All modules and Team Access"],
], [1600, 4140, 3620])
add_callout(doc, "Access principle", "Navigation reduces clutter, but backend role and factory guards remain the enforcement point. Never share credentials or reuse another person’s account.", fill="E7EFE8")

# 9. Supporting modules
add_heading(doc, "9. Supporting modules", page_break=True)
add_table(doc, ["Module", "Use", "Control"], [
    ["Expenses", "Record factory and vehicle costs; allocate eligible cost to raw blocks", "Manager/owner review"],
    ["Machines", "Maintain B-21/LPM register and daily runtime logs", "Manager/owner administration"],
    ["Opening Setup", "Guided physical count before go-live", "Approval required before factory becomes live"],
    ["Tally Imports", "Upload daybook or trial-balance XML as separate import batches", "Manager/owner only; separate from live inventory"],
    ["Historical Data", "Backfill historical sales summaries with reason and audit context", "Manager/owner only"],
    ["Team Access", "Provision a signed-up user into the factory with a role", "Owner controls owner grants"],
], [1900, 4660, 2800])
add_heading(doc, "Opening a new factory", 2)
add_steps(doc, [
    "Create the opening count and enter raw blocks, unpolished slabs and finished slabs.",
    "Resolve ambiguous or incomplete lines before submission.",
    "Submit the physical count for approval.",
    "Approve the snapshot; StoneOS appends opening-receipt movements and marks the factory live.",
])
add_callout(doc, "Administrative separation", "Operational users should not use imports or historical backfill to fix today’s production, inventory or sales records.")

# 10. Daily playbooks
add_heading(doc, "10. Daily playbooks", page_break=True)
add_heading(doc, "Operator", 2)
add_bullets(doc, ["Confirm the correct block, slab batch and machine before starting.", "Enter runtime and downtime against the operational date.", "Complete the event only after the physical step is finished.", "Escalate eligibility warnings instead of bypassing them."])
add_heading(doc, "Supervisor", 2)
add_bullets(doc, ["Review My Work at shift start and after major completions.", "Verify B-21 good/damaged counts and LPM hand-offs.", "Resolve material-location mismatches before completion or abort.", "Review active sessions at shift handover."])
add_heading(doc, "Manager or owner", 2)
add_bullets(doc, ["Review exceptions, reversals, open reservations and partial dispatches.", "Check daily production and commercial summaries.", "Approve user access and opening-stock actions deliberately.", "Keep historical imports and live operations separate."])
add_heading(doc, "End-of-shift checklist", 2)
add_bullets(doc, ["No physically completed run remains open without a reason.", "No slab is shown in a location different from its physical rack or WIP area.", "Reserved slabs match open customer commitments.", "Dispatches and invoices posted today reconcile to the intended loads."])

# 11. Exceptions and troubleshooting
add_heading(doc, "11. Exceptions and troubleshooting", page_break=True)
add_table(doc, ["Message or symptom", "Meaning", "What to do"], [
    ["Not eligible", "Stage, location, status or reservation does not match", "Open the slab/block record; reconcile the physical state before retrying"],
    ["Changed after session started", "Reserved WIP was adjusted or moved", "Do not complete/abort; ask a manager to reconcile inventory and the active reservation"],
    ["Unable to reach API", "Frontend cannot contact the backend", "Confirm the backend health endpoint, API URL and local/network connectivity"],
    ["API error 500", "The backend failed while processing the request", "Do not repeat mutations rapidly; capture the action and time, then check backend logs"],
    ["Already reversed", "A correction movement already exists", "Review movement history; never create a second compensating reversal"],
    ["Unauthorized", "Role, factory mapping or session is invalid", "Sign in again; owner/manager should verify Team Access and factory assignment"],
], [2500, 3000, 3860])
add_callout(doc, "When reporting an issue", "Include the page, action, serial/order ID, exact time and visible message. Do not send passwords, authentication tokens or live customer financial data in screenshots.", fill="FFF6D9", accent=BRASS)

# 12. Quick reference
add_heading(doc, "12. Quick reference and glossary", page_break=True)
add_table(doc, ["Term", "Definition"], [
    ["Operational date", "The production day used for logs and summaries; configured around the factory shift boundary."],
    ["B-21 WIP", "The location/state for a raw block actively being cut."],
    ["LPM Queue", "The controlled queue between grinding, epoxy and polishing events."],
    ["Reservation", "A temporary hold that prevents the same material from being used by another workflow."],
    ["Movement", "An append-only inventory ledger event recording source, destination, reason and actor."],
    ["Reversal", "A linked compensating movement used to correct an earlier event without deleting history."],
    ["Genealogy", "The trace from slab to parent block and production sessions, then forward to customer activity."],
    ["Tenant", "One factory’s isolated data boundary."],
], [2100, 7260])
add_heading(doc, "Demo dataset used in this guide", 2)
add_bullets(doc, [
    "Factory: Atlas Stone Works (fictional).",
    "Blocks: V101–V104 and slabs derived from those demo serials.",
    "Customer: Northstar Surfaces (fictional).",
    "Orders, invoices, dispatches, quantities and rupee values are training examples only.",
])
add_callout(doc, "Final reminder", "Record the physical event once, in the correct workflow, at the time it happens. StoneOS will carry the material identity and audit trail forward.", fill="E7EFE8")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(OUT)
